from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_report():
    doc = Document()

    # Styles setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('RAPPORT – IMPLÉMENTATION DU SÉLECTEUR DE DENTS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('(Niveau technique – équipe de développement)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph('_' * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER # Horizontal line
    doc.add_paragraph()

    # 1. Contexte
    doc.add_heading('1. Contexte et objectif', level=1)
    doc.add_paragraph(
        "Dans le cadre de l’évolution de l’application médicale, il est nécessaire d’implémenter un sélecteur de dents destiné aux utilisateurs ayant le rôle de dentiste."
    )
    doc.add_paragraph(
        "Ce composant permettra d’identifier précisément les dents concernées lors d’une consultation, de les associer à des actes médicaux et de transmettre ces informations aux modules de facturation."
    )
    p = doc.add_paragraph()
    runner = p.add_run("Objectif principal : ")
    runner.bold = True
    p.add_run("Créer un composant modulaire, maintenable et indépendant, facilement intégrable dans plusieurs parties de l’application.")

    # 2. Prérequis
    doc.add_heading('2. Prérequis', level=1)
    doc.add_paragraph("Avant de démarrer l’implémentation, les prérequis suivants doivent être satisfaits :")
    
    doc.add_heading('2.1 Prérequis techniques', level=2)
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Application basée sur React (≥ 17, idéalement 18)")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Structure de projet déjà découpée en : pages, components, hooks")

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Système de rôles utilisateurs existant : Dentiste, Médecin, Secrétaire, Administrateur")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Étape 1 (nettoyage et découpage des gros fichiers) au moins partiellement réalisée")

    doc.add_heading('2.2 Prérequis fonctionnels', level=2)
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Existence d’une consultation dentaire")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Compréhension de la norme FDI (numérotation internationale des dents)")
    
    # Placeholder Image FDI
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    p_img = cell.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run("\n[IMAGE SUGGÉRÉE : SCHÉMA DE LA NUMÉROTATION FDI]\n(Insérer ici un visuel de la mâchoire numérotée 11-48)\n")
    run_img.italic = True
    run_img.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Accès aux modules : consultation, actes médicaux, facturation")

    # 3. Objectif fonctionnel
    doc.add_heading('3. Objectif fonctionnel du sélecteur', level=1)
    doc.add_paragraph("Le sélecteur de dents doit permettre au dentiste de :")
    
    items = [
        "Sélectionner une ou plusieurs dents",
        "Définir l’état de chaque dent (ex. : cariée, extraite)",
        "Visualiser les dents sélectionnées en mode lecture seule",
        "Transmettre une structure de données exploitable par les modules métiers"
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    warning = doc.add_paragraph()
    run_warn = warning.add_run("⚠️ Le composant ne doit pas :")
    run_warn.bold = True
    run_warn.font.color.rgb = RGBColor(200, 0, 0)
    
    negatives = [
        "Connaître le patient",
        "Gérer la facturation",
        "Appeler directement une API",
        "Dépendre du contexte global de l’application"
    ]
    for item in negatives:
        doc.add_paragraph(item, style='List Bullet')

    # 4. Choix techniques
    doc.add_heading('4. Choix techniques retenus', level=1)
    
    doc.add_heading('4.1 Rendu graphique', level=2)
    doc.add_paragraph("• Utilisation d’un SVG interactif représentant la mâchoire dentaire")
    doc.add_paragraph("• Chaque dent est identifiée par un ID conforme à la norme FDI (11 à 48)")
    
    doc.add_paragraph("Justification : Précision médicale, légèreté, facilité de maintenance, compatibilité responsive.")

    doc.add_heading('4.2 Gestion de la logique', level=2)
    doc.add_paragraph("• La logique de sélection est isolée dans un hook dédié")
    doc.add_paragraph("• Le composant principal se limite à l’affichage et aux interactions")

    # 5. Modules
    doc.add_heading('5. Modules et dépendances npm', level=1)
    doc.add_paragraph("Les dépendances suivantes sont recommandées :")
    
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.5)
    run_code = p_code.add_run("npm install clsx")
    run_code.font.name = 'Courier New'
    run_code.font.color.rgb = RGBColor(0, 50, 150)
    run_code.element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")

    doc.add_paragraph("Détails :")
    doc.add_paragraph("• clsx : gestion propre des classes CSS selon l’état des dents, remplace les conditions CSS complexes", style='List Bullet')
    doc.add_paragraph("Les autres dépendances (React, lucide-react) sont déjà présentes dans le projet.")
    
    p = doc.add_paragraph()
    run = p.add_run("❌ Aucune librairie lourde (canvas, d3, three.js) n’est nécessaire.")
    run.bold = True

    # 6. Structure
    doc.add_heading('6. Structure recommandée des fichiers', level=1)
    doc.add_paragraph("Le composant doit être isolé dans un dossier dédié :")

    structure_text = """components/
└── dentist/
    └── ToothSelector/
        ├── ToothSelector.jsx      (composant principal)
        ├── Tooth.jsx              (dent individuelle)
        ├── ToothLegend.jsx        (légende des états)
        ├── ToothMap.js            (cartographie FDI)
        ├── constants.js           (états des dents)
        └── hooks/
            └── useToothSelector.js"""
    
    # Add code block with border
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    p_struct = cell.paragraphs[0]
    run_struct = p_struct.add_run(structure_text)
    run_struct.font.name = 'Courier New'
    run_struct.font.size = Pt(9)
    
    doc.add_paragraph("Cette structure permet : une maintenance facilitée, une évolution future (implant, prothèse) et une intégration simple.")

    # 7. Modèle de données
    doc.add_heading('7. Modèle de données manipulé', level=1)
    doc.add_paragraph("Le sélecteur doit produire une structure stable (ex: liste des dents sélectionnées + état). Ces données seront stockées avec la consultation et exploitées par la facturation.")
    
    p = doc.add_paragraph()
    run = p.add_run("⚠️ Ce format doit être validé une fois et ne plus changer.")
    run.bold = True

    # 8. Étapes
    doc.add_heading('8. Étapes de développement', level=1)
    steps = [
        "1. Définition de la cartographie des dents (norme FDI)",
        "2. Création du hook de gestion de sélection",
        "3. Développement du composant “dent individuelle”",
        "4. Assemblage dans le composant principal",
        "5. Ajout de la légende et des états",
        "6. Tests fonctionnels en environnement isolé"
    ]
    for step in steps:
        doc.add_paragraph(step)

    # 9. Intégration
    doc.add_heading('9. Intégration dans l’application', level=1)
    doc.add_paragraph("L’intégration doit être progressive :")
    integ_steps = [
        "Test du composant dans une page dédiée",
        "Intégration dans la consultation dentaire",
        "Restriction de l’affichage au rôle dentiste",
        "Connexion aux actes médicaux",
        "Exploitation dans le module de facturation"
    ]
    for s in integ_steps:
        doc.add_paragraph(s, style='List Number')

    # 10. Résultat
    doc.add_heading('10. Résultat attendu', level=1)
    doc.add_paragraph("À la fin de cette étape :")
    res_items = [
        "Le sélecteur de dents est fonctionnel",
        "Le code est modulaire et maintenable",
        "Le composant est prêt à évoluer",
        "La base est posée pour le module de caisse"
    ]
    for item in res_items:
        doc.add_paragraph(item, style='List Bullet')

    # Save
    file_path = "/mnt/data/Rapport_Implementation_Selecteur_Dents.docx"
    doc.save(file_path)
    return file_path

file_path = create_report()
print(file_path)